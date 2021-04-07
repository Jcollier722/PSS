from docx import Document
from docx.shared import Inches
from datetime import datetime

def export(path,sjn,fcfs,srt,rr):

    #create and title new document
    document = Document()
    document.add_heading('CPU Scheduling Simulation Results', 0)

    #add timestamp of generation
    sub_title = document.add_paragraph('Generated '+str(datetime.now().strftime("%d/%m/%Y %H:%M:%S")))

    #make the table of averages
    avg_table(document,sjn,fcfs,srt,rr)

    #write fcfs and sjn, then page-break (so tables dont wrap pages)
    write_algo(document,"First Come, First Serve",fcfs)
    write_algo(document,"Shortest Job Next",sjn)
    document.add_page_break()
    #finish up with srt and rr 
    write_algo(document,"Shortest Remaining Time",srt)
    write_algo(document,"Round Robin",rr)
    
    #close and save doc
    document.save(path)

"""Make the table for all of the averages -> fixed amount of cells here since we are working with 4 algos """
def avg_table(document,sjn,fcfs,srt,rr):
    document.add_heading('Average across all algorithms', level=2)
    table = document.add_table(rows=5, cols=3)
    table.style='Light Shading'
    header = table.rows[0].cells
    
    header[0].text = "Algorithm"
    header[1].text = "Average Turnaround Time"
    header[2].text = "Average Wait Time"

    row_1 = table.rows[1].cells
    row_1[0].text= "FCFS"
    row_1[1].text= str(fcfs[1])
    row_1[2].text = str(fcfs[2])

    row_2 = table.rows[2].cells
    row_2[0].text= "SJN"
    row_2[1].text= str(sjn[1])
    row_2[2].text = str(sjn[2])

    row_3 = table.rows[3].cells
    row_3[0].text= "SRT"
    row_3[1].text= str(srt[1])
    row_3[2].text = str(srt[2])

    row_4 = table.rows[4].cells
    row_4[0].text= "RR"
    row_4[1].text= str(rr[1])
    row_4[2].text = str(rr[2])

""" Write the results of the simulation into a table -> pass the current document, the name of the alogirthm and the results of that simulation"""
def write_algo(document,algoname,results):
    document.add_heading(algoname,level=2)
    table = document.add_table(rows=1, cols=3)
    table.style='Light Shading'
    header = table.rows[0].cells
    
    header[0].text = "Job"
    header[1].text = "Turnaround Time"
    header[2].text = "Wait Time"

    job_list = results[0]

    for job in job_list:
        this_row = table.add_row().cells
        this_row[0].text= str(job.name)
        this_row[1].text= str(job.turn)
        this_row[2].text= str(job.wait)
    


        
